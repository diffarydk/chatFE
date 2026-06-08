from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "deliverables"
DOCX = OUT / "Hands-On_TUBES_Implementasi_Chat_Real-Time_Cloud.docx"

NAVY = "17365D"
BLUE = "2E75B6"
DARK = "222222"
GRAY = "666666"
WHITE = "FFFFFF"
PALE = "F2F5F8"
LIGHT_BLUE = "D9EAF7"
GREEN = "E2F0D9"
YELLOW = "FFF2CC"
RED = "FCE4D6"
WIDTH = 9360


def font(
    run,
    name="Calibri",
    size=10.5,
    color=DARK,
    bold=False,
    italic=False,
):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), name)
    rpr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def shade(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    node = tcpr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tcpr.append(node)
    node.set(qn("w:fill"), fill)


def cell_margin(cell, top=80, start=120, bottom=80, end=120):
    tcpr = cell._tc.get_or_add_tcPr()
    margins = tcpr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tcpr.append(margins)

    for side, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width):
    tcpr = cell._tc.get_or_add_tcPr()
    tcw = tcpr.find(qn("w:tcW"))
    if tcw is None:
        tcw = OxmlElement("w:tcW")
        tcpr.append(tcw)
    tcw.set(qn("w:w"), str(width))
    tcw.set(qn("w:type"), "dxa")
    cell_margin(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table_geometry(table, widths, indent=120):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tblpr = table._tbl.tblPr

    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(sum(widths)))
    tblw.set(qn("w:type"), "dxa")

    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tblpr.append(tblind)
    tblind.set(qn("w:w"), str(indent))
    tblind.set(qn("w:type"), "dxa")

    layout = tblpr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblpr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        trpr = row._tr.get_or_add_trPr()
        trpr.append(OxmlElement("w:cantSplit"))
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])


def page_field(paragraph):
    run = paragraph.add_run("Halaman ")
    font(run, size=8.5, color=GRAY)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, text, end])


def setup(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.32)
    section.footer_distance = Inches(0.32)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.widow_control = True

    for name, size, before, after in (
        ("Heading 1", 15, 13, 7),
        ("Heading 2", 12.5, 10, 5),
        ("Heading 3", 11, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(BLUE)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.space_after = Pt(1)
    run = header.add_run(
        "HANDS-ON TUGAS BESAR  |  CLOUD COMPUTING  |  BECHAT + FECHAT"
    )
    font(run, size=8, color=GRAY, bold=True)
    ppr = header._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), BLUE)
    pbdr.append(bottom)
    ppr.append(pbdr)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    page_field(footer)


class Lists:
    def __init__(self, doc):
        self.doc = doc

    def new(self, kind="number"):
        return "List Bullet" if kind == "bullet" else "List Number"

    def item(self, text, style):
        paragraph = self.doc.add_paragraph(style=style)
        font(paragraph.add_run(text))
        return paragraph


def body(doc, text, bold_prefix=None, italic=False, justify=False):
    paragraph = doc.add_paragraph()
    if justify:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_prefix and text.startswith(bold_prefix):
        font(paragraph.add_run(bold_prefix), bold=True)
        font(paragraph.add_run(text[len(bold_prefix) :]), italic=italic)
    else:
        font(paragraph.add_run(text), italic=italic)
    return paragraph


def band(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table_geometry(table, [WIDTH], indent=0)
    cell = table.cell(0, 0)
    shade(cell, NAVY)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    font(paragraph.add_run(text.upper()), size=13, color=WHITE, bold=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def heading(doc, text, level=2):
    return doc.add_paragraph(text, style=f"Heading {level}")


def code(doc, value, label=None):
    if label:
        label_paragraph = doc.add_paragraph()
        label_paragraph.paragraph_format.keep_with_next = True
        label_paragraph.paragraph_format.space_after = Pt(2)
        font(
            label_paragraph.add_run(label),
            size=9.5,
            color=NAVY,
            bold=True,
        )

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.04)
    paragraph.paragraph_format.right_indent = Inches(0.04)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1
    paragraph.paragraph_format.keep_together = True

    ppr = paragraph._p.get_or_add_pPr()
    background = OxmlElement("w:shd")
    background.set(qn("w:fill"), PALE)
    ppr.append(background)

    borders = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        edge = OxmlElement(f"w:{side}")
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "4")
        edge.set(qn("w:space"), "4")
        edge.set(qn("w:color"), "C9D2DC")
        borders.append(edge)
    ppr.append(borders)

    font(paragraph.add_run(value), name="Courier New", size=8.1)
    return paragraph


def note(doc, title, text, kind="info"):
    fill = {
        "info": LIGHT_BLUE,
        "warn": YELLOW,
        "ok": GREEN,
        "risk": RED,
    }[kind]
    table = doc.add_table(rows=1, cols=1)
    table_geometry(table, [WIDTH])
    cell = table.cell(0, 0)
    shade(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    font(paragraph.add_run(title), size=10, color=NAVY, bold=True)
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    font(paragraph.add_run(text), size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def data_table(doc, headers, rows, widths, size=8.6):
    table = doc.add_table(rows=1, cols=len(headers))
    table_geometry(table, widths)

    for index, value in enumerate(headers):
        cell = table.rows[0].cells[index]
        shade(cell, NAVY)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        font(paragraph.add_run(value), size=size, color=WHITE, bold=True)

    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            set_cell_width(cells[index], widths[index])
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            font(paragraph.add_run(str(value)), size=size)

    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_steps(lists, items, kind="number"):
    style = lists.new(kind)
    for item in items:
        lists.item(item, style)


def cover(doc):
    doc.add_paragraph().paragraph_format.space_after = Pt(42)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(
        paragraph.add_run("HANDS-ON TUGAS BESAR"),
        size=14,
        color=GRAY,
        bold=True,
    )

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(5)
    font(
        paragraph.add_run(
            "Implementasi Aplikasi Chat Real-Time\n"
            "Berbasis Cloud Menggunakan Docker dan AWS"
        ),
        size=25,
        color=NAVY,
        bold=True,
    )

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(25)
    font(
        paragraph.add_run("Studi Kasus: Backend BECHAT dan Frontend FECHAT"),
        size=14,
        color=BLUE,
        bold=True,
    )

    data_table(
        doc,
        ["Aspek", "Deskripsi"],
        [
            ("Jenis Dokumen", "Dokumentasi hands-on implementasi Tugas Besar"),
            ("Aplikasi", "Chat real-time cloud-native BECHAT dan FECHAT"),
            ("Backend", "FastAPI, Socket.IO, SQLAlchemy Async, JWT, bcrypt"),
            ("Frontend", "React 19, Vite, TypeScript, Tailwind CSS, React Router"),
            ("Database", "SQLite development dan PostgreSQL 15 container"),
            ("Container", "Docker multi-stage, Docker Compose, dan Nginx"),
            ("Cloud", "Amazon ECR/Docker Hub dan deployment AWS EC2/ECS"),
            (
                "Repository",
                "BECHAT: github.com/diffarydk/BECHAT | "
                "FECHAT: github.com/diffarydk/chatFE",
            ),
            ("Tahun", "2026"),
        ],
        [2250, 7110],
        9.1,
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(20)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(paragraph.add_run("CLOUD COMPUTING"), size=13, color=NAVY, bold=True)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(
        paragraph.add_run("IMPLEMENTASI FULL-STACK DAN DEPLOYMENT CLOUD"),
        size=11,
        color=GRAY,
        bold=True,
    )


def backend_hands_on(doc, lists):
    band(doc, "Pendahuluan")
    body(
        doc,
        "Dokumen ini menjelaskan proses implementasi aplikasi chat real-time "
        "secara full-stack. Bagian I membahas backend BECHAT dan melanjutkan "
        "hasil hands-on backend yang telah diuji pada 5 Juni 2026. Bagian II "
        "membahas frontend FECHAT dan cara menghubungkannya ke kontrak REST API "
        "serta Socket.IO backend.",
        justify=True,
    )
    body(
        doc,
        "BECHAT menangani autentikasi, database, REST API, file sharing, dan "
        "event real-time. FECHAT menjadi client React yang menyimpan session, "
        "melindungi route privat, memanggil API, menampilkan workspace chat, "
        "dan dapat diaktifkan untuk menerima pesan Socket.IO tanpa refresh.",
        justify=True,
    )
    heading(doc, "Teknologi yang Digunakan")
    data_table(
        doc,
        ["Lapisan", "Teknologi", "Fungsi"],
        [
            ("Backend API", "FastAPI + Uvicorn", "REST API berbasis ASGI"),
            ("Real-time", "python-socketio", "Event pesan dan status pengguna"),
            ("Database", "SQLAlchemy Async", "ORM SQLite/PostgreSQL"),
            ("Security", "JWT + bcrypt", "Token dan hashing password"),
            ("Frontend", "React + Vite", "Single-page application"),
            ("UI", "Tailwind CSS + Motion", "Layout dan animasi"),
            ("Container", "Docker + Nginx", "Runtime konsisten"),
            ("Cloud", "AWS + registry image", "Distribusi dan deployment"),
        ],
        [1450, 2600, 5310],
    )
    heading(doc, "Alur Implementasi Full-Stack")
    code(
        doc,
        "Frontend FECHAT :3000 / :80\n"
        "        |\n"
        "        | REST API + Bearer JWT\n"
        "        | Socket.IO + auth token\n"
        "        v\n"
        "Backend BECHAT :5000\n"
        "        |\n"
        "        +-> PostgreSQL :5432\n"
        "        +-> uploads / Amazon S3\n\n"
        "GitHub -> CI/CD -> Image Registry -> AWS EC2/ECS",
    )

    band(doc, "Bagian I - Hands-On Backend BECHAT")

    band(doc, "1. Persiapan Environment")
    body(
        doc,
        "Implementasi backend dilakukan menggunakan Python, Docker Desktop, "
        "Docker Compose, dan Git. Pastikan seluruh aplikasi pendukung tersedia."
    )
    add_steps(
        lists,
        [
            "Periksa versi Python, Git, Docker, dan Docker Compose.",
            "Masuk ke repository backend D:\\BECHAT.",
            "Periksa branch aktif dan remote repository sebelum mengubah kode.",
        ],
    )
    code(
        doc,
        "python --version\n"
        "git --version\n"
        "docker --version\n"
        "docker compose version",
        "Verifikasi aplikasi pendukung",
    )
    code(
        doc,
        "cd D:\\BECHAT\n"
        "git remote -v\n"
        "git branch --show-current",
        "Masuk ke repository",
    )

    band(doc, "2. Struktur Project BECHAT")
    body(
        doc,
        "Repository backend disusun modular agar konfigurasi, model, schema, "
        "autentikasi, endpoint, dan data seed tidak berada dalam satu file."
    )
    code(
        doc,
        "BECHAT/\n"
        "|-- .github/workflows/docker-be.yml\n"
        "|-- app/\n"
        "|   |-- main.py\n"
        "|   |-- config.py\n"
        "|   |-- database.py\n"
        "|   |-- models.py\n"
        "|   |-- schemas.py\n"
        "|   |-- auth.py\n"
        "|   `-- seed.py\n"
        "|-- Dockerfile\n"
        "|-- docker-compose.yml\n"
        "|-- requirements.txt\n"
        "|-- .env.example\n"
        "`-- BACKEND_INTEGRATION.md",
    )
    data_table(
        doc,
        ["File", "Isi Implementasi"],
        [
            ("app/main.py", "FastAPI, REST endpoint, Socket.IO, lifecycle"),
            ("app/config.py", "Environment variable dan Pydantic Settings"),
            ("app/database.py", "Async engine dan session SQLAlchemy"),
            ("app/models.py", "Struktur tabel dan relationship"),
            ("app/schemas.py", "Validasi request dan response"),
            ("app/auth.py", "bcrypt, JWT, dan HTTPBearer"),
            ("app/seed.py", "Data awal untuk pengujian"),
        ],
        [2500, 6860],
    )

    band(doc, "3. Arsitektur Backend dan Cloud")
    body(
        doc,
        "Backend menggabungkan FastAPI dan Socket.IO pada satu ASGI application "
        "agar REST API dan koneksi real-time memakai port 5000 yang sama."
    )
    code(
        doc,
        "Browser / Frontend FECHAT\n"
        "        |\n"
        "        | REST API + Bearer JWT\n"
        "        | Socket.IO + auth token\n"
        "        v\n"
        "FastAPI + Socket.IO ASGI :5000\n"
        "        |\n"
        "        +-> SQLAlchemy Async -> PostgreSQL :5432\n"
        "        +-> uploads/ volume -> file sharing",
    )
    code(
        doc,
        'app = FastAPI(title=settings.APP_NAME, lifespan=lifespan)\n\n'
        'sio = socketio.AsyncServer(async_mode="asgi")\n'
        "sio_asgi_app = socketio.ASGIApp(\n"
        "    socketio_server=sio,\n"
        "    other_asgi_app=app,\n"
        '    socketio_path="/socket.io",\n'
        ")",
        "Pembuatan ASGI application",
    )

    band(doc, "4. Konfigurasi Environment")
    code(
        doc,
        "JWT_SECRET=your_jwt_secret_key_here\n"
        "DATABASE_URL=postgresql+psycopg://username:password@host:port/database\n"
        'CORS_ORIGINS=["http://localhost:3000","http://localhost:5173"]',
        ".env.example",
    )
    code(
        doc,
        "JWT_SECRET=ganti-dengan-random-secret-yang-panjang\n"
        "DATABASE_URL=sqlite+aiosqlite:///./app.db\n"
        'CORS_ORIGINS=["http://localhost:3000"]',
        "Development dengan SQLite",
    )
    note(
        doc,
        "Perhatian credential",
        "Jangan menyimpan JWT secret, password database, access key AWS, atau "
        "private key .pem di repository Git.",
        "risk",
    )

    band(doc, "5. Implementasi Struktur Database")
    data_table(
        doc,
        ["Tabel", "Kolom Utama", "Kegunaan"],
        [
            ("users", "id, name, email, password_hash, status", "Akun/presence"),
            ("chats", "id, name, type, last_message", "Direct/group chat"),
            ("chat_members", "chat_id, user_id, role", "Relasi user-chat"),
            ("messages", "chat_id, sender_id, content, status", "Riwayat pesan"),
            ("friend_requests", "from_id, to_id, status", "Pertemanan"),
            ("files", "owner_id, chat_id, name, url", "Metadata file"),
        ],
        [1850, 3500, 4010],
    )
    code(
        doc,
        'class Message(Base):\n'
        '    __tablename__ = "messages"\n'
        "    id = Column(String, primary_key=True)\n"
        '    chat_id = Column(String, ForeignKey("chats.id"))\n'
        '    sender_id = Column(String, ForeignKey("users.id"))\n'
        "    content = Column(String, nullable=False)\n"
        '    status = Column(String, default="sent")',
        "Contoh model Message",
    )

    band(doc, "6. Implementasi Autentikasi dan Security")
    body(
        doc,
        "Password disimpan sebagai hash bcrypt. Setelah login berhasil, backend "
        "membuat JWT yang dikirim frontend melalui Authorization Bearer."
    )
    code(
        doc,
        "def hash_password(password: str) -> str:\n"
        '    password_bytes = password.encode("utf-8")\n'
        "    return bcrypt.hashpw(password_bytes, bcrypt.gensalt()).decode()\n\n"
        "def verify_password(plain: str, hashed: str) -> bool:\n"
        '    return bcrypt.checkpw(plain.encode("utf-8"), hashed.encode("utf-8"))',
        "Hashing password",
    )
    add_steps(
        lists,
        [
            "Frontend mengirim email dan password ke POST /api/auth/login.",
            "Backend mencari user dan memverifikasi password dengan bcrypt.",
            "Backend membuat JWT jika credential benar.",
            "Frontend menyimpan token dan mengirim Bearer token pada API privat.",
        ],
    )

    band(doc, "7. Implementasi REST API")
    data_table(
        doc,
        ["Method", "Endpoint", "Fungsi"],
        [
            ("POST", "/api/auth/register", "Membuat akun"),
            ("POST", "/api/auth/login", "Login dan JWT"),
            ("GET", "/api/chats", "Daftar chat"),
            ("GET", "/api/chats/{id}/messages", "Riwayat pesan"),
            ("GET", "/api/members", "Direktori member"),
            ("GET", "/api/friend-requests", "Request pending"),
            ("POST", "/api/friend-requests", "Mengirim request"),
            ("POST", "/api/friend-requests/{id}/accept", "Menerima request"),
            ("POST", "/api/friend-requests/{id}/reject", "Menolak request"),
            ("GET/POST", "/api/files", "Daftar dan upload file"),
        ],
        [1150, 4100, 4110],
        8.2,
    )
    code(
        doc,
        "POST http://localhost:5000/api/auth/login\n"
        "Content-Type: application/json\n\n"
        '{\n  "email": "ada@example.com",\n  "password": "Secret123"\n}',
        "Contoh request login",
    )

    band(doc, "8. Implementasi Chat Real-Time dengan Socket.IO")
    code(
        doc,
        "const socket = io('http://localhost:5000', {\n"
        "  auth: { token: localStorage.getItem('token') }\n"
        "});\n\n"
        "socket.emit('joinChat', { chatId: 'c1' });\n"
        "socket.emit('sendMessage', {\n"
        "  chatId: 'c1',\n"
        "  content: 'Halo dari frontend'\n"
        "});",
        "Integrasi client",
    )
    data_table(
        doc,
        ["Arah", "Event", "Proses"],
        [
            ("Client -> Server", "joinChat", "Masuk room chat"),
            ("Client -> Server", "leaveChat", "Keluar room"),
            ("Client -> Server", "sendMessage", "Simpan/broadcast pesan"),
            ("Server -> Client", "receiveMessage", "Pesan baru"),
            ("Server -> Client", "chatUpdated", "Update sidebar chat"),
            ("Server -> Client", "memberStatusUpdate", "Presence"),
            ("Server -> Client", "friendRequestReceived", "Notifikasi request"),
        ],
        [1900, 2600, 4860],
        8.4,
    )

    band(doc, "9. Implementasi Friend Request dan File Sharing")
    body(
        doc,
        "Ketika friend request diterima, backend membuat direct chat dan "
        "menambahkan kedua user sebagai anggota. File diunggah memakai "
        "multipart/form-data; binary disimpan pada volume dan metadata pada DB."
    )
    note(
        doc,
        "Penyimpanan production",
        "Local volume sesuai untuk satu instance. Deployment multi-instance "
        "sebaiknya menggunakan Amazon S3 dan signed URL.",
    )

    band(doc, "10. Data Seeding")
    data_table(
        doc,
        ["User ID", "Nama", "Email", "Password Uji", "Role"],
        [
            ("ada-77", "Ada Lovelace", "ada@example.com", "Secret123", "Admin"),
            ("bob-99", "Bob Engineer", "bob@example.com", "Secret123", "Engineer"),
            ("elena", "Elena Rostova", "elena@example.com", "Secret123", "Admin"),
            (
                "marcus",
                "Marcus Chen",
                "marcus@example.com",
                "Secret123",
                "Engineer",
            ),
        ],
        [1350, 2050, 2700, 1800, 1460],
        8.5,
    )
    note(
        doc,
        "Khusus development",
        "Credential seed bersifat publik dan harus diganti atau dinonaktifkan "
        "pada production.",
        "warn",
    )

    band(doc, "11. Menjalankan Backend Secara Lokal")
    code(
        doc,
        "cd D:\\BECHAT\n"
        "python -m venv .venv\n"
        ".\\.venv\\Scripts\\Activate.ps1\n"
        "pip install -r requirements.txt\n"
        "python -m uvicorn app.main:sio_asgi_app --reload --port 5000",
        "Mode SQLite",
    )
    data_table(
        doc,
        ["Layanan", "Alamat"],
        [
            ("Root/health", "http://localhost:5000/"),
            ("Swagger UI", "http://localhost:5000/docs"),
            ("OpenAPI JSON", "http://localhost:5000/openapi.json"),
            ("Socket.IO", "http://localhost:5000/socket.io"),
        ],
        [2400, 6960],
    )

    band(doc, "12. Implementasi Docker")
    code(
        doc,
        "FROM python:3.13-slim as builder\n"
        "WORKDIR /app\n"
        "RUN python -m venv /opt/venv\n"
        "COPY requirements.txt .\n"
        "RUN /opt/venv/bin/pip install -r requirements.txt\n\n"
        "FROM python:3.13-slim as runner\n"
        "WORKDIR /app\n"
        "COPY --from=builder /opt/venv /opt/venv\n"
        "COPY app /app/app\n"
        "EXPOSE 5000\n"
        'CMD ["uvicorn", "app.main:sio_asgi_app", '
        '"--host", "0.0.0.0", "--port", "5000"]',
        "Ringkasan Dockerfile",
    )
    code(
        doc,
        "docker compose up --build -d\n"
        "docker compose ps\n"
        "docker compose logs -f backend",
        "Menjalankan container",
    )

    band(doc, "13. Hasil Pengujian Docker dan Database")
    body(
        doc,
        "Pengujian backend menunjukkan container backend dan PostgreSQL healthy, "
        "seed berhasil, dan enam tabel utama terbentuk."
    )
    code(
        doc,
        "public | chat_members    | table\n"
        "public | chats           | table\n"
        "public | files           | table\n"
        "public | friend_requests | table\n"
        "public | messages        | table\n"
        "public | users           | table",
        "Verifikasi tabel PostgreSQL",
    )

    band(doc, "14. Hasil Pengujian REST API")
    note(
        doc,
        "Status pengujian",
        "POST /api/auth/login, GET /api/chats, GET /api/members, dan "
        "GET /api/friend-requests menghasilkan HTTP 200.",
        "ok",
    )

    band(doc, "15. Hasil Pengujian Socket.IO")
    note(
        doc,
        "Hasil",
        "Client terhubung menggunakan JWT, bergabung ke chat c1, dan menerima "
        "event receiveMessage tanpa refresh halaman.",
        "ok",
    )

    band(doc, "16. Version Control dan GitHub")
    code(
        doc,
        "git status\n"
        "git add app Dockerfile docker-compose.yml requirements.txt\n"
        'git commit -m "feat: implement real-time cloud chat backend"\n'
        "git push origin main",
    )
    code(
        doc,
        ".env\n.venv/\n__pycache__/\n*.db\nuploads/\n*.pem",
        "File yang tidak boleh di-push",
    )

    band(doc, "17. GitHub Actions dan Amazon ECR")
    add_steps(
        lists,
        [
            "Checkout source code.",
            "Menyiapkan Docker Buildx.",
            "Login ke Amazon ECR menggunakan credential pipeline.",
            "Build image backend.",
            "Push tag latest dan tag berdasarkan Git commit SHA.",
        ],
    )
    note(
        doc,
        "Rekomendasi",
        "Gunakan GitHub OpenID Connect dan IAM Role untuk menghindari "
        "long-lived access key.",
    )

    band(doc, "18. Deployment Backend ke AWS")
    data_table(
        doc,
        ["Port", "Source", "Kegunaan"],
        [
            ("22", "IP administrator", "SSH"),
            ("80", "0.0.0.0/0", "HTTP reverse proxy"),
            ("443", "0.0.0.0/0", "HTTPS dan WSS"),
            ("5000", "Terbatas/internal", "Debug backend"),
        ],
        [1300, 2900, 5160],
    )
    code(
        doc,
        "sudo apt update\n"
        "sudo apt install docker.io docker-compose-plugin awscli -y\n"
        "sudo usermod -aG docker $USER\n"
        "newgrp docker",
        "Install Docker pada EC2",
    )
    code(
        doc,
        "aws ecr get-login-password --region us-east-1 | \\\n"
        "docker login --username AWS --password-stdin \\\n"
        "058264530865.dkr.ecr.us-east-1.amazonaws.com\n\n"
        "docker pull "
        "058264530865.dkr.ecr.us-east-1.amazonaws.com/fechat-be:latest",
        "Login dan pull image",
    )

    band(doc, "19. Domain, HTTPS, dan WebSocket")
    body(
        doc,
        "Production membutuhkan HTTPS untuk REST API dan WSS untuk Socket.IO. "
        "Reverse proxy atau load balancer harus meneruskan upgrade WebSocket."
    )
    code(
        doc,
        "server {\n"
        "    listen 443 ssl;\n"
        "    server_name api.example.com;\n\n"
        "    location / {\n"
        "        proxy_pass http://127.0.0.1:5000;\n"
        "        proxy_http_version 1.1;\n"
        "        proxy_set_header Upgrade $http_upgrade;\n"
        '        proxy_set_header Connection "upgrade";\n'
        "        proxy_set_header Host $host;\n"
        "    }\n"
        "}",
        "Contoh konfigurasi Nginx",
    )

    band(doc, "20. Security Hardening")
    add_steps(
        lists,
        [
            "Ganti JWT_SECRET dan simpan pada secret manager.",
            "Batasi CORS hanya untuk domain frontend yang sah.",
            "Jangan expose PostgreSQL ke internet.",
            "Gunakan RDS private subnet dan migration Alembic.",
            "Pindahkan file upload ke S3 serta validasi ukuran dan MIME.",
            "Tambahkan rate limiting login/register.",
            "Gunakan image tag SHA dan aktifkan monitoring.",
            "Rotasi credential atau private key yang pernah tersimpan di Git.",
        ],
        "bullet",
    )

    band(doc, "21. Troubleshooting Backend")
    data_table(
        doc,
        ["Masalah", "Penyebab dan Penyelesaian"],
        [
            ("Docker engine tidak ditemukan", "Jalankan Docker Desktop."),
            ("Port 5000 digunakan", "Hentikan proses lama atau ubah port."),
            ("PostgreSQL refused", "Gunakan hostname postgres di Compose."),
            ("401 Unauthorized", "Kirim Bearer token yang valid."),
            ("Socket rejected", "Pastikan auth.token dan JWT valid."),
            ("Seed tidak berjalan", "Seed dilewati jika users sudah terisi."),
            ("WebSocket gagal HTTPS", "Teruskan header Upgrade."),
        ],
        [2600, 6760],
        8.5,
    )

    band(doc, "22. Ringkasan Implementasi Backend")
    body(
        doc,
        "Backend BECHAT telah menyediakan REST API, autentikasi JWT, database "
        "asynchronous, Socket.IO, Docker Compose, dan distribusi image. Bagian "
        "berikut melanjutkan hands-on dengan implementasi frontend FECHAT.",
        justify=True,
    )
    note(
        doc,
        "Titik Lanjut",
        "Pastikan backend dapat diakses pada http://localhost:5000 dan CORS "
        "mengizinkan http://localhost:3000 sebelum memulai Bagian II.",
        "ok",
    )


def frontend_hands_on(doc, lists):
    doc.add_page_break()
    band(doc, "Bagian II - Hands-On Frontend FECHAT")
    body(
        doc,
        "Bagian ini melanjutkan hands-on backend dengan repository FECHAT pada "
        "D:\\FECHAT. Seluruh contoh mengikuti source code aktual. Fitur login, "
        "register, route guard, API wrapper, initial chat data, user lookup, dan "
        "friend request telah memiliki titik integrasi backend. Socket.IO dan "
        "beberapa halaman lain masih perlu diaktifkan dari implementasi mock.",
        justify=True,
    )

    band(doc, "23. Persiapan Environment Frontend")
    add_steps(
        lists,
        [
            "Pastikan Node.js 20 atau versi LTS tersedia.",
            "Masuk ke repository D:\\FECHAT.",
            "Install dependency dari package-lock.json.",
            "Buat file .env.local untuk alamat backend.",
            "Jalankan Vite pada port 3000.",
        ],
    )
    code(
        doc,
        "node --version\nnpm --version\n\n"
        "cd D:\\FECHAT\n"
        "npm install\n"
        "npm run dev",
        "Menjalankan FECHAT",
    )
    note(
        doc,
        "Alamat development",
        "Script dev menjalankan Vite pada http://localhost:3000 dan dapat "
        "diakses dari host lain karena memakai --host=0.0.0.0.",
    )

    band(doc, "24. Struktur Project FECHAT")
    code(
        doc,
        "FECHAT/\n"
        "|-- src/\n"
        "|   |-- components/\n"
        "|   |   |-- ProtectedRoute.tsx\n"
        "|   |   `-- PublicRoute.tsx\n"
        "|   |-- context/AuthContext.tsx\n"
        "|   |-- pages/\n"
        "|   |   |-- Login.tsx\n"
        "|   |   |-- Register.tsx\n"
        "|   |   |-- Chat.tsx\n"
        "|   |   |-- GroupChats.tsx\n"
        "|   |   |-- Files.tsx\n"
        "|   |   `-- Members.tsx\n"
        "|   |-- services/api.ts\n"
        "|   |-- App.tsx\n"
        "|   |-- main.tsx\n"
        "|   `-- index.css\n"
        "|-- Dockerfile\n"
        "|-- nginx.conf\n"
        "|-- vite.config.ts\n"
        "|-- package.json\n"
        "`-- .env.example",
    )
    data_table(
        doc,
        ["File", "Peran"],
        [
            ("src/App.tsx", "Router dan daftar route aplikasi"),
            ("src/context/AuthContext.tsx", "State user/token dan localStorage"),
            ("src/services/api.ts", "Wrapper fetch dan Bearer token"),
            ("src/pages/Login.tsx", "Login real atau fallback mock"),
            ("src/pages/Register.tsx", "Register real atau fallback mock"),
            ("src/pages/Chat.tsx", "Chat, member, request, integrasi API"),
            ("Dockerfile", "Build Vite lalu serve melalui Nginx"),
            ("nginx.conf", "SPA fallback ke index.html"),
        ],
        [3000, 6360],
    )

    band(doc, "25. Konfigurasi Environment FECHAT")
    code(
        doc,
        "VITE_API_URL=http://localhost:5000/api\n"
        "VITE_WS_URL=http://localhost:5000",
        ".env.local",
    )
    body(
        doc,
        "VITE_API_URL wajib diisi untuk memakai backend nyata. Halaman Login, "
        "Register, dan Chat memeriksa environment ini; jika kosong, aplikasi "
        "menggunakan mock data untuk demonstrasi offline."
    )
    note(
        doc,
        "Perilaku Vite",
        "Environment variable berawalan VITE_ dimasukkan ke bundle saat proses "
        "build. Mengubah URL pada server setelah image dibangun tidak mengubah "
        "bundle; image harus dibangun ulang.",
        "warn",
    )

    band(doc, "26. Routing dan Route Guard")
    code(
        doc,
        "<AuthProvider>\n"
        "  <BrowserRouter>\n"
        "    <Routes>\n"
        '      <Route path="/login" element={<PublicRoute><Login /></PublicRoute>} />\n'
        '      <Route path="/register" element={<PublicRoute><Register /></PublicRoute>} />\n'
        '      <Route path="/chat" element={<ProtectedRoute><Chat /></ProtectedRoute>} />\n'
        '      <Route path="/groups" element={<ProtectedRoute><GroupChats /></ProtectedRoute>} />\n'
        '      <Route path="/files" element={<ProtectedRoute><Files /></ProtectedRoute>} />\n'
        '      <Route path="/members" element={<ProtectedRoute><Members /></ProtectedRoute>} />\n'
        "    </Routes>\n"
        "  </BrowserRouter>\n"
        "</AuthProvider>",
        "Struktur App.tsx",
    )
    data_table(
        doc,
        ["Route", "Jenis", "Perilaku"],
        [
            ("/login", "Publik", "Redirect ke /chat jika sudah login"),
            ("/register", "Publik", "Redirect ke /chat jika sudah login"),
            ("/chat", "Privat", "Redirect ke /login tanpa token"),
            ("/groups", "Privat", "Halaman group chat"),
            ("/files", "Privat", "Halaman file"),
            ("/members", "Privat", "Direktori member"),
        ],
        [1800, 1800, 5760],
    )

    band(doc, "27. AuthContext dan Penyimpanan Session")
    body(
        doc,
        "AuthProvider membaca token dan user dari localStorage ketika aplikasi "
        "dimulai. Route guard menunggu isLoading selesai sebelum memutuskan "
        "redirect."
    )
    code(
        doc,
        "const login = (newToken: string, newUser: User) => {\n"
        "  setToken(newToken);\n"
        "  setUser(newUser);\n"
        "  localStorage.setItem('token', newToken);\n"
        "  localStorage.setItem('user', JSON.stringify(newUser));\n"
        "};\n\n"
        "const logout = () => {\n"
        "  setToken(null);\n"
        "  setUser(null);\n"
        "  localStorage.removeItem('token');\n"
        "  localStorage.removeItem('user');\n"
        "};",
        "Session frontend",
    )
    note(
        doc,
        "Catatan security",
        "localStorage sederhana untuk hands-on, tetapi rentan jika aplikasi "
        "memiliki celah XSS. Production dapat memakai access token singkat dan "
        "refresh token pada cookie httpOnly.",
        "warn",
    )

    band(doc, "28. API Wrapper dan Bearer Token")
    code(
        doc,
        "const BASE_URL = import.meta.env.VITE_API_URL ||\n"
        "  'http://localhost:5000/api';\n\n"
        "const token = localStorage.getItem('token');\n"
        "if (token) {\n"
        "  headers.set('Authorization', `Bearer ${token}`);\n"
        "}\n\n"
        "const response = await fetch(`${BASE_URL}${endpoint}`, config);",
        "src/services/api.ts",
    )
    add_steps(
        lists,
        [
            "Wrapper mengambil token dari localStorage.",
            "Authorization Bearer ditambahkan otomatis.",
            "Content-Type JSON ditambahkan kecuali body berupa FormData.",
            "Response non-2xx diubah menjadi Error.",
            "Status 401 menghapus session dan redirect ke /login.",
        ],
    )
    note(
        doc,
        "Kontrak response",
        "Wrapper selalu memanggil response.json(). Endpoint sukses sebaiknya "
        "mengembalikan JSON. Download file/blob memerlukan fetch khusus.",
    )

    band(doc, "29. Integrasi Login dan Register")
    code(
        doc,
        "const data = await api.post('/auth/login', {\n"
        "  email,\n"
        "  password\n"
        "});\n"
        "login(data.token, data.user);\n"
        "navigate('/chat');",
        "Login.tsx",
    )
    code(
        doc,
        "const data = await api.post('/auth/register', {\n"
        "  userId,\n"
        "  name,\n"
        "  email,\n"
        "  password\n"
        "});\n"
        "login(data.token, data.user);\n"
        "navigate('/chat');",
        "Register.tsx",
    )
    data_table(
        doc,
        ["Endpoint", "Request", "Response minimum"],
        [
            (
                "POST /auth/login",
                "email, password",
                "token, user{id,name,email,avatar}",
            ),
            (
                "POST /auth/register",
                "userId, name, email, password",
                "token, user{id,name,email,avatar}",
            ),
        ],
        [2600, 3000, 3760],
    )
    note(
        doc,
        "Pemetaan ID",
        "Register mengirim userId, tetapi frontend menyimpan user.id. Backend "
        "harus mengembalikan field id pada object user.",
        "warn",
    )

    band(doc, "30. Memuat Chat, Member, dan Friend Request")
    code(
        doc,
        "const chatsData = await api.get('/chats');\n"
        "setChats(chatsData);\n\n"
        "const membersData = await api.get('/members');\n"
        "setMembers(membersData);\n\n"
        "const requestData = await api.get('/friend-requests');\n"
        "setFriendRequests(requestData);",
        "Initial data pada Chat.tsx",
    )
    data_table(
        doc,
        ["Endpoint", "Dipakai untuk"],
        [
            ("GET /chats", "Sidebar direct dan group chat"),
            ("GET /members", "Panel member workspace"),
            ("GET /friend-requests", "Modal friend request"),
            ("GET /users/{id}", "Validasi user saat memulai chat"),
            ("POST /friend-requests", "Mengirim friend request"),
        ],
        [3600, 5760],
    )
    code(
        doc,
        "const user = await api.get(`/users/${targetId}`);\n"
        "await api.post('/friend-requests', { to: targetId });\n"
        "alert(`Friend request sent to ${user.name}`);",
        "Mencari user dan mengirim friend request",
    )

    band(doc, "31. Mengaktifkan Message History")
    body(
        doc,
        "Source saat ini belum mengambil riwayat pesan ketika activeChat "
        "berubah. Tambahkan effect berikut setelah backend menyediakan endpoint "
        "GET /api/chats/{chatId}/messages."
    )
    code(
        doc,
        "useEffect(() => {\n"
        "  if (!activeChat || !import.meta.env.VITE_API_URL) return;\n\n"
        "  api.get(`/chats/${activeChat}/messages`)\n"
        "    .then(setMessages)\n"
        "    .catch((error) => console.error(error));\n"
        "}, [activeChat]);",
        "Tambahan pada Chat.tsx",
    )
    note(
        doc,
        "Format response",
        "Jika backend mengembalikan object pagination {items,nextCursor}, ubah "
        "setMessages(data) menjadi setMessages(data.items).",
    )

    band(doc, "32. Mengaktifkan Socket.IO pada FECHAT")
    body(
        doc,
        "Chat.tsx saat ini hanya berisi placeholder Socket.IO dan simulasi "
        "status sent/delivered/read. Install client lalu buat satu koneksi yang "
        "memakai token login."
    )
    code(doc, "npm install socket.io-client", "Install dependency")
    code(
        doc,
        "import { io } from 'socket.io-client';\n\n"
        "const WS_URL = import.meta.env.VITE_WS_URL ||\n"
        "  'http://localhost:5000';\n\n"
        "export const socket = io(WS_URL, {\n"
        "  autoConnect: false,\n"
        "  auth: { token: localStorage.getItem('token') }\n"
        "});",
        "src/services/socket.ts",
    )
    code(
        doc,
        "useEffect(() => {\n"
        "  if (!token) return;\n"
        "  socket.auth = { token };\n"
        "  socket.connect();\n\n"
        "  socket.on('receiveMessage', (message) => {\n"
        "    setMessages((current) => [...current, message]);\n"
        "  });\n\n"
        "  socket.on('messageStatusUpdate', ({ messageId, status }) => {\n"
        "    setMessages((current) => current.map((message) =>\n"
        "      message.id === messageId ? { ...message, status } : message\n"
        "    ));\n"
        "  });\n\n"
        "  return () => {\n"
        "    socket.removeAllListeners();\n"
        "    socket.disconnect();\n"
        "  };\n"
        "}, [token]);",
        "Connect dan listen event",
    )
    code(
        doc,
        "socket.emit('sendMessage', {\n"
        "  chatId: activeChat,\n"
        "  content: inputValue\n"
        "});",
        "Mengirim pesan",
    )
    note(
        doc,
        "Hapus simulasi",
        "Setelah Socket.IO aktif, hapus setTimeout yang mengubah status pesan "
        "secara lokal. Status harus berasal dari event backend.",
        "warn",
    )

    band(doc, "33. Join dan Leave Room Chat")
    code(
        doc,
        "useEffect(() => {\n"
        "  if (!activeChat || !socket.connected) return;\n"
        "  socket.emit('joinChat', { chatId: activeChat });\n\n"
        "  return () => {\n"
        "    socket.emit('leaveChat', { chatId: activeChat });\n"
        "  };\n"
        "}, [activeChat]);",
        "Sinkronisasi room",
    )
    body(
        doc,
        "Backend harus mengabaikan senderId dari payload client dan memakai "
        "identitas user hasil verifikasi JWT pada Socket.IO session."
    )

    band(doc, "34. Menyelesaikan Friend Request")
    body(
        doc,
        "Handler accept dan reject pada Chat.tsx masih mengubah state lokal. "
        "Aktifkan endpoint backend agar perubahan persisten."
    )
    code(
        doc,
        "const result = await api.post(\n"
        "  `/friend-requests/${requestId}/accept`\n"
        ");\n"
        "setFriendRequests((current) =>\n"
        "  current.filter((request) => request.id !== requestId)\n"
        ");\n"
        "setChats((current) => [result.chat, ...current]);",
        "Accept request",
    )
    code(
        doc,
        "await api.post(`/friend-requests/${requestId}/reject`);\n"
        "setFriendRequests((current) =>\n"
        "  current.filter((request) => request.id !== requestId)\n"
        ");",
        "Reject request",
    )

    band(doc, "35. Status Halaman Group, Files, dan Members")
    data_table(
        doc,
        ["Halaman", "Kondisi Saat Ini", "Langkah Integrasi"],
        [
            ("GroupChats.tsx", "Data dan pesan mock", "GET /groups dan Socket.IO"),
            ("Files.tsx", "Daftar file mock", "GET/POST/DELETE /files"),
            ("Members.tsx", "Direktori member mock", "GET /members"),
        ],
        [2400, 3000, 3960],
    )
    note(
        doc,
        "Urutan yang disarankan",
        "Stabilkan login, register, chat list, message history, dan Socket.IO "
        "lebih dahulu. Setelah itu pindahkan GroupChats, Files, dan Members dari "
        "mock data ke backend.",
    )

    band(doc, "36. Pengujian Full-Stack Lokal")
    add_steps(
        lists,
        [
            "Jalankan backend pada http://localhost:5000.",
            "Pastikan CORS backend mengizinkan http://localhost:3000.",
            "Isi VITE_API_URL dan VITE_WS_URL pada .env.local.",
            "Jalankan npm run dev.",
            "Register atau login memakai akun seed.",
            "Periksa Network tab untuk status REST API.",
            "Periksa koneksi Socket.IO dan kirim pesan dari dua browser.",
        ],
    )
    code(
        doc,
        "npm run lint\nnpm run build",
        "Validasi source frontend sebelum deployment",
    )
    data_table(
        doc,
        ["Skenario", "Hasil yang Diharapkan"],
        [
            ("Login valid", "Redirect ke /chat dan token tersimpan"),
            ("Refresh /chat", "Session dipulihkan dari localStorage"),
            ("Token invalid", "Session dihapus dan redirect /login"),
            ("GET /chats", "Sidebar menampilkan data backend"),
            ("Kirim pesan", "Browser lain menerima tanpa refresh"),
            ("Accept request", "Direct chat baru muncul"),
        ],
        [3000, 6360],
    )

    band(doc, "37. Dockerfile Frontend dan Nginx")
    code(
        doc,
        "FROM node:20-alpine AS builder\n"
        "WORKDIR /app\n"
        "COPY package*.json ./\n"
        "RUN npm install\n"
        "COPY . .\n"
        "ARG VITE_API_URL\n"
        "ARG VITE_WS_URL\n"
        "ENV VITE_API_URL=$VITE_API_URL\n"
        "ENV VITE_WS_URL=$VITE_WS_URL\n"
        "RUN npm run build\n\n"
        "FROM nginx:alpine\n"
        "RUN rm /etc/nginx/conf.d/default.conf\n"
        "COPY nginx.conf /etc/nginx/conf.d/default.conf\n"
        "COPY --from=builder /app/dist /usr/share/nginx/html\n"
        "EXPOSE 80\n"
        'CMD ["nginx", "-g", "daemon off;"]',
        "Dockerfile production",
    )
    code(
        doc,
        "server {\n"
        "    listen 80;\n"
        "    server_name localhost;\n\n"
        "    location / {\n"
        "        root /usr/share/nginx/html;\n"
        "        index index.html;\n"
        "        try_files $uri $uri/ /index.html;\n"
        "    }\n"
        "}",
        "SPA fallback pada nginx.conf",
    )

    band(doc, "38. Build dan Menjalankan Image Frontend")
    code(
        doc,
        "docker build \\\n"
        "  --build-arg VITE_API_URL=https://api.example.com/api \\\n"
        "  --build-arg VITE_WS_URL=https://api.example.com \\\n"
        "  -t USERNAME/chat-frontend:latest .\n\n"
        "docker run -d \\\n"
        "  --name chat-frontend \\\n"
        "  --restart unless-stopped \\\n"
        "  -p 80:80 \\\n"
        "  USERNAME/chat-frontend:latest",
    )
    note(
        doc,
        "Build-time configuration",
        "VITE_API_URL harus berisi prefix /api, sedangkan VITE_WS_URL menunjuk "
        "origin Socket.IO tanpa suffix /api.",
        "warn",
    )

    band(doc, "39. Deployment Frontend ke AWS EC2")
    add_steps(
        lists,
        [
            "Buat EC2 Ubuntu dan buka port 22, 80, dan 443.",
            "Install Docker pada instance.",
            "Pull image frontend dari Docker Hub atau Amazon ECR.",
            "Jalankan container frontend pada port 80.",
            "Arahkan DNS frontend ke Elastic IP EC2.",
            "Pasang HTTPS menggunakan ACM/ALB atau Certbot.",
        ],
    )
    code(
        doc,
        "docker pull USERNAME/chat-frontend:latest\n"
        "docker stop chat-frontend || true\n"
        "docker rm chat-frontend || true\n"
        "docker run -d \\\n"
        "  --name chat-frontend \\\n"
        "  -p 80:80 \\\n"
        "  --restart always \\\n"
        "  USERNAME/chat-frontend:latest",
        "Deployment container",
    )

    band(doc, "40. Domain, HTTPS, CORS, dan WSS")
    body(
        doc,
        "Contoh topologi production menggunakan https://chat.example.com untuk "
        "frontend dan https://api.example.com untuk backend. Browser memanggil "
        "REST melalui HTTPS dan Socket.IO melalui WSS."
    )
    data_table(
        doc,
        ["Komponen", "Nilai Production"],
        [
            ("Frontend", "https://chat.example.com"),
            ("VITE_API_URL", "https://api.example.com/api"),
            ("VITE_WS_URL", "https://api.example.com"),
            ("Backend CORS", '["https://chat.example.com"]'),
            ("Protocol realtime", "WSS melalui port 443"),
        ],
        [3000, 6360],
    )
    note(
        doc,
        "Mixed content",
        "Frontend HTTPS tidak boleh memanggil API atau Socket.IO melalui HTTP. "
        "Gunakan HTTPS/WSS pada kedua sisi.",
        "risk",
    )

    band(doc, "41. Troubleshooting Frontend")
    data_table(
        doc,
        ["Masalah", "Penyebab dan Penyelesaian"],
        [
            ("Frontend tetap memakai mock", "Isi VITE_API_URL lalu restart Vite."),
            ("Failed to fetch", "Periksa URL, backend, CORS, dan firewall."),
            ("401 lalu kembali login", "Token invalid/expired; login ulang."),
            ("Route 404 setelah refresh", "Pastikan try_files ke index.html."),
            ("Socket tidak connect", "Periksa VITE_WS_URL, token, path, dan WSS."),
            ("Pesan tampil dua kali", "Hindari listener ganda dan cleanup effect."),
            ("URL production tidak berubah", "Build ulang image dengan build-arg."),
            ("Download file gagal", "Gunakan fetch blob, bukan api.get JSON."),
        ],
        [2700, 6660],
        8.4,
    )

    band(doc, "42. Checklist Integrasi End-to-End")
    add_steps(
        lists,
        [
            "Backend health check dan Swagger dapat diakses.",
            "CORS backend mengizinkan origin frontend.",
            "Login/register mengembalikan token dan user.id.",
            "Route privat hanya terbuka setelah autentikasi.",
            "Chat, member, dan friend request berasal dari backend.",
            "Message history dimuat saat activeChat berubah.",
            "Socket.IO mengirim token pada handshake.",
            "joinChat, leaveChat, sendMessage, dan receiveMessage berfungsi.",
            "Simulasi status pesan lokal sudah dihapus.",
            "Image frontend dibangun dengan API/WS URL production.",
            "Frontend dan backend memakai HTTPS/WSS.",
            "Secret dan private key tidak tersimpan di Git.",
        ],
        "bullet",
    )

    band(doc, "43. Kesimpulan Full-Stack")
    body(
        doc,
        "Hands-on backend BECHAT telah dilanjutkan dengan frontend FECHAT. "
        "Frontend menggunakan React, Vite, TypeScript, route guard, AuthContext, "
        "dan API wrapper untuk mengonsumsi layanan backend secara aman.",
        justify=True,
    )
    body(
        doc,
        "Integrasi dasar REST telah memiliki titik implementasi langsung pada "
        "Login, Register, dan Chat. Message history, Socket.IO, accept/reject "
        "friend request, serta halaman Group, Files, dan Members dijelaskan "
        "sebagai tahap implementasi lanjutan berdasarkan kondisi source aktual.",
        justify=True,
    )
    body(
        doc,
        "Docker multi-stage membangun bundle Vite dan Nginx menyajikan SPA. "
        "Dengan URL API/WS yang benar, CORS terbatas, HTTPS/WSS, image registry, "
        "dan AWS EC2/ECS, sistem dapat dijalankan sebagai aplikasi chat "
        "full-stack berbasis cloud.",
        justify=True,
    )
    data_table(
        doc,
        ["Komponen", "Status dalam Repository"],
        [
            ("React/Vite/Tailwind", "Tersedia"),
            ("Routing dan route guard", "Tersedia"),
            ("AuthContext/localStorage", "Tersedia"),
            ("REST API wrapper", "Tersedia"),
            ("Login/register backend", "Titik integrasi tersedia"),
            ("Initial chat data", "Titik integrasi tersedia"),
            ("Socket.IO client", "Perlu diaktifkan"),
            ("Group/Files/Members API", "Masih mock"),
            ("Docker + Nginx SPA", "Tersedia"),
            ("Deployment EC2", "Panduan dan script tersedia"),
        ],
        [4900, 4460],
        9,
    )
    note(
        doc,
        "Akhir Hands-On",
        "Dokumen menggabungkan implementasi backend BECHAT dan kelanjutan "
        "frontend FECHAT berdasarkan source code repository masing-masing.",
        "ok",
    )


def references(doc, lists):
    band(doc, "Referensi")
    add_steps(
        lists,
        [
            "Dokumentasi FastAPI - https://fastapi.tiangolo.com",
            "Dokumentasi Socket.IO - https://socket.io/docs/v4",
            "Dokumentasi SQLAlchemy - https://docs.sqlalchemy.org",
            "Dokumentasi React - https://react.dev",
            "Dokumentasi Vite - https://vite.dev",
            "Dokumentasi React Router - https://reactrouter.com",
            "Dokumentasi Tailwind CSS - https://tailwindcss.com",
            "Dokumentasi Docker - https://docs.docker.com",
            "Dokumentasi Nginx - https://nginx.org/en/docs",
            "Dokumentasi Amazon EC2 - https://docs.aws.amazon.com/ec2",
            "Dokumentasi Amazon ECR - https://docs.aws.amazon.com/ecr",
            "Repository BECHAT - https://github.com/diffarydk/BECHAT",
            "Repository FECHAT - https://github.com/diffarydk/chatFE",
        ],
    )


def build():
    OUT.mkdir(exist_ok=True)
    doc = Document()
    setup(doc)
    lists = Lists(doc)
    cover(doc)
    doc.add_page_break()

    backend_hands_on(doc, lists)
    frontend_hands_on(doc, lists)
    references(doc, lists)

    doc.core_properties.title = (
        "Implementasi Aplikasi Chat Real-Time Berbasis Cloud Menggunakan "
        "Docker dan AWS"
    )
    doc.core_properties.subject = (
        "Hands-On Tugas Besar Backend BECHAT dan Frontend FECHAT"
    )
    doc.core_properties.author = (
        "Dokumentasi berdasarkan repository BECHAT dan FECHAT"
    )
    doc.core_properties.keywords = (
        "BECHAT, FECHAT, FastAPI, React, Socket.IO, Docker, PostgreSQL, AWS"
    )
    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    build()
